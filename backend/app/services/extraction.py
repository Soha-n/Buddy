"""Turn an uploaded file into retrievable text chunks.

Every parser here returns the same shape - a list of (text, locator) pairs -
so the ingestion pipeline downstream never branches on file type. The locator
is a human-readable origin ("page 4", "Sheet1 rows 20-40") kept alongside the
text so an answer can say where a passage came from instead of citing an
opaque chunk id.

Tabular files (CSV/Excel) are treated differently from prose on purpose. Text
chunks of a spreadsheet retrieve badly - a question like "what were Q3 sales"
matches the header row, not the number - so alongside the chunks they also
produce a *data summary*: real column names, dtypes and a small preview, which
is always injected into the prompt rather than retrieved. That is what lets the
model write pandas code against actual column names instead of guessing them.
"""

from __future__ import annotations

import csv
import io
import logging
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

# Chunk sizing is in characters, not tokens: an exact token count would need the
# model's own tokenizer, and these bounds only have to be roughly right. ~1200
# chars lands near 300 tokens, small enough that several chunks fit in a modest
# context window alongside the conversation.
CHUNK_CHARS = 1200
CHUNK_OVERLAP_CHARS = 180

# A hard ceiling on extracted text per file. Long documents still work - they
# just stop contributing new chunks past this point, which keeps a 500-page PDF
# from producing an indexing job that appears to hang.
MAX_EXTRACT_CHARS = 600_000

# Rows/cols read from a spreadsheet. Wide enough for real reports, bounded so a
# million-row export cannot exhaust memory during ingestion.
MAX_TABLE_ROWS = 5_000
MAX_TABLE_COLS = 64
PREVIEW_ROWS = 8

# Values pandas renders for missing data; skipped when serializing rows so a
# sparse sheet does not fill chunks with "nan".
_EMPTY_CELL_TOKENS = {"", "nan", "NaN", "None", "<NA>", "NaT"}


class ExtractionError(Exception):
    """The file could not be read. Message is shown to the user verbatim."""


@dataclass
class Chunk:
    content: str
    locator: str | None = None


@dataclass
class Extracted:
    chunks: list[Chunk]
    #: Set for tabular files only; always injected rather than retrieved.
    data_summary: str | None = None
    truncated: bool = False


# --------------------------------------------------------------------------- #
# Kind detection
# --------------------------------------------------------------------------- #

_TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".json", ".xml", ".yaml", ".yml",
}
_TABLE_SUFFIXES = {".csv", ".tsv", ".xlsx", ".xlsm", ".xls"}
_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}


def detect_kind(filename: str) -> str:
    """Classify a file into pdf / docx / table / text / image / unsupported.

    Extension-based rather than content-sniffing: these are files the user
    picked from their own disk, and a wrong guess from magic bytes would be
    harder to explain than "we don't support .xyz".
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix in {".docx", ".docm"}:
        return "docx"
    if suffix in _TABLE_SUFFIXES:
        return "table"
    if suffix in _TEXT_SUFFIXES:
        return "text"
    if suffix in _IMAGE_SUFFIXES:
        return "image"
    return "unsupported"


# --------------------------------------------------------------------------- #
# Chunking
# --------------------------------------------------------------------------- #


def chunk_text(text: str, locator: str | None = None) -> list[Chunk]:
    """Split text on paragraph boundaries, packing up to CHUNK_CHARS.

    Paragraph-first (rather than a blind character slice) keeps sentences and
    list items intact, because a chunk cut mid-sentence embeds toward the wrong
    meaning. Consecutive chunks overlap by CHUNK_OVERLAP_CHARS so a fact sitting
    exactly on a boundary is still fully present in at least one chunk.
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        return []

    paragraphs = [p.strip() for p in normalized.split("\n\n") if p.strip()]
    if not paragraphs:
        paragraphs = [normalized]

    chunks: list[Chunk] = []
    buffer = ""

    def flush() -> None:
        nonlocal buffer
        if buffer.strip():
            chunks.append(Chunk(content=buffer.strip(), locator=locator))
        buffer = ""

    for paragraph in paragraphs:
        # A single paragraph longer than the budget is hard-split; there is
        # nothing smaller to break on.
        if len(paragraph) > CHUNK_CHARS:
            flush()
            step = CHUNK_CHARS - CHUNK_OVERLAP_CHARS
            for start in range(0, len(paragraph), step):
                piece = paragraph[start : start + CHUNK_CHARS].strip()
                if piece:
                    chunks.append(Chunk(content=piece, locator=locator))
            continue

        if len(buffer) + len(paragraph) + 2 > CHUNK_CHARS:
            flush()
            # Carry the tail of the previous chunk forward as overlap.
            if chunks and CHUNK_OVERLAP_CHARS > 0:
                buffer = chunks[-1].content[-CHUNK_OVERLAP_CHARS:] + "\n\n"
        buffer += paragraph + "\n\n"

    flush()
    return chunks


# --------------------------------------------------------------------------- #
# Per-format extraction
# --------------------------------------------------------------------------- #


def _extract_pdf(data: bytes) -> Extracted:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise ExtractionError("PDF support is not installed on the backend.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"This PDF could not be opened: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # An empty password unlocks a surprising number of "protected" PDFs.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ExtractionError(
                "This PDF is password-protected. Remove the password and re-upload."
            ) from exc

    chunks: list[Chunk] = []
    total_chars = 0
    truncated = False

    for page_number, page in enumerate(reader.pages, start=1):
        if total_chars >= MAX_EXTRACT_CHARS:
            truncated = True
            break
        try:
            page_text = page.extract_text() or ""
        except Exception:
            # One malformed page should not sink the whole document.
            logger.debug("failed to extract PDF page %d", page_number, exc_info=True)
            continue
        if not page_text.strip():
            continue
        total_chars += len(page_text)
        chunks.extend(chunk_text(page_text, locator=f"page {page_number}"))

    if not chunks:
        raise ExtractionError(
            "No text could be read from this PDF. It looks like a scanned image, "
            "and Buddy cannot OCR it yet."
        )
    return Extracted(chunks=chunks, truncated=truncated)


def _extract_docx(data: bytes) -> Extracted:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("DOCX support is not installed on the backend.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(
            f"This Word file could not be opened: {exc}. "
            "Legacy .doc files must be saved as .docx first."
        ) from exc

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Tables carry real content in most reports, so they are flattened into
    # pipe-delimited rows rather than dropped.
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {table_index}]\n" + "\n".join(rows))

    text = "\n\n".join(parts)
    if not text.strip():
        raise ExtractionError("This document appears to be empty.")

    truncated = len(text) > MAX_EXTRACT_CHARS
    return Extracted(chunks=chunk_text(text[:MAX_EXTRACT_CHARS]), truncated=truncated)


def _extract_plain_text(data: bytes) -> Extracted:
    # utf-8 first, then cp1252 (the usual Windows source of stray bytes), then a
    # lossy pass so one bad byte never fails an otherwise readable file.
    text: str | None = None
    for encoding in ("utf-8", "cp1252"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = data.decode("utf-8", errors="replace")

    if not text.strip():
        raise ExtractionError("This file appears to be empty.")

    truncated = len(text) > MAX_EXTRACT_CHARS
    return Extracted(chunks=chunk_text(text[:MAX_EXTRACT_CHARS]), truncated=truncated)


def _sniff_csv_delimiter(sample: str, filename: str) -> str:
    if Path(filename).suffix.lower() == ".tsv":
        return "\t"
    try:
        return csv.Sniffer().sniff(sample, delimiters=",;\t|").delimiter
    except csv.Error:
        return ","


def _frame_summary(frame, source_label: str) -> str:
    """Column names, dtypes and a preview - the model's map of the data.

    Always injected rather than retrieved: without real column names the model
    invents them, and generated pandas code then fails on a KeyError.
    """
    lines = [
        f"{source_label}: {len(frame)} rows x {len(frame.columns)} columns",
        "Columns (name :: dtype):",
    ]
    for column in frame.columns:
        lines.append(f"  - {column} :: {frame[column].dtype}")

    # A head-only preview is actively misleading: a model reading "First 8 rows"
    # of a 9-row file treats it as the whole file and answers max/min questions
    # from a truncated view. Showing the tail as well, and labelling the preview
    # as partial, keeps it from being mistaken for the complete data.
    try:
        if len(frame) <= PREVIEW_ROWS * 2:
            preview = frame.to_string(max_cols=MAX_TABLE_COLS)
            label = f"All {len(frame)} rows"
        else:
            head = frame.head(PREVIEW_ROWS).to_string(max_cols=MAX_TABLE_COLS)
            tail = frame.tail(PREVIEW_ROWS).to_string(
                max_cols=MAX_TABLE_COLS, header=False
            )
            omitted = len(frame) - PREVIEW_ROWS * 2
            preview = f"{head}\n   ... {omitted} more rows ...\n{tail}"
            label = (
                f"Sample only - first and last {PREVIEW_ROWS} of {len(frame)} rows. "
                f"Write code to read every row"
            )
    except Exception:
        preview = "(preview unavailable)"
        label = "Preview"
    lines.append(f"{label}:\n{preview}")

    # describe() is what turns "what's the average" into a directly answerable
    # question with no code execution at all.
    numeric = frame.select_dtypes(include="number")
    if not numeric.empty:
        try:
            lines.append(f"Numeric summary:\n{numeric.describe().to_string()}")
        except Exception:
            pass

    return "\n".join(lines)


def _chunks_from_frame(frame, sheet_label: str) -> list[Chunk]:
    """Serialize rows as 'column: value' records, batched into chunks.

    Row-wise labelled records rather than a raw CSV dump: each retrieved chunk
    then carries its own column names, so a chunk from the middle of a sheet is
    still interpretable without the header row.
    """
    chunks: list[Chunk] = []
    buffer: list[str] = []
    buffer_chars = 0
    start_row = 0

    for position, (_, row) in enumerate(frame.iterrows()):
        record = "; ".join(
            f"{column}: {row[column]}"
            for column in frame.columns
            if str(row[column]).strip() not in _EMPTY_CELL_TOKENS
        )
        if not record:
            continue

        if buffer_chars + len(record) > CHUNK_CHARS and buffer:
            chunks.append(
                Chunk(
                    content="\n".join(buffer),
                    locator=f"{sheet_label} {start_row + 1}-{position}",
                )
            )
            buffer, buffer_chars, start_row = [], 0, position

        buffer.append(record)
        buffer_chars += len(record) + 1

    if buffer:
        chunks.append(
            Chunk(
                content="\n".join(buffer),
                locator=f"{sheet_label} {start_row + 1}-{len(frame)}",
            )
        )
    return chunks


def _extract_table(data: bytes, filename: str) -> Extracted:
    try:
        import pandas as pd
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "Spreadsheet support is not installed on the backend."
        ) from exc

    suffix = Path(filename).suffix.lower()
    chunks: list[Chunk] = []
    summaries: list[str] = []
    truncated = False

    if suffix in {".csv", ".tsv"}:
        sample = data[:64_000].decode("utf-8", errors="replace")
        delimiter = _sniff_csv_delimiter(sample, filename)
        try:
            frame = pd.read_csv(
                io.BytesIO(data),
                delimiter=delimiter,
                nrows=MAX_TABLE_ROWS,
                encoding_errors="replace",
                on_bad_lines="skip",
            )
        except Exception as exc:
            raise ExtractionError(f"This CSV could not be parsed: {exc}") from exc

        if len(frame.columns) > MAX_TABLE_COLS:
            frame = frame.iloc[:, :MAX_TABLE_COLS]
            truncated = True
        if len(frame) >= MAX_TABLE_ROWS:
            truncated = True

        summaries.append(_frame_summary(frame, filename))
        chunks.extend(_chunks_from_frame(frame, "rows"))
    else:
        try:
            sheets = pd.read_excel(
                io.BytesIO(data), sheet_name=None, nrows=MAX_TABLE_ROWS
            )
        except Exception as exc:
            raise ExtractionError(
                f"This spreadsheet could not be opened: {exc}. "
                "Legacy .xls files may need re-saving as .xlsx."
            ) from exc

        for sheet_name, frame in sheets.items():
            if frame.empty:
                continue
            if len(frame.columns) > MAX_TABLE_COLS:
                frame = frame.iloc[:, :MAX_TABLE_COLS]
                truncated = True
            if len(frame) >= MAX_TABLE_ROWS:
                truncated = True
            summaries.append(_frame_summary(frame, f"Sheet '{sheet_name}'"))
            chunks.extend(_chunks_from_frame(frame, f"Sheet '{sheet_name}' rows"))

    if not chunks and not summaries:
        raise ExtractionError("No data rows were found in this file.")

    return Extracted(
        chunks=chunks,
        data_summary="\n\n".join(summaries) if summaries else None,
        truncated=truncated,
    )


def extract(data: bytes, filename: str) -> Extracted:
    """Dispatch to the parser for this file type."""
    kind = detect_kind(filename)
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "table":
        return _extract_table(data, filename)
    if kind == "text":
        return _extract_plain_text(data)
    if kind == "image":
        # Images carry no extractable text; their searchable content is the
        # vision-model description generated after upload.
        return Extracted(chunks=[])
    raise ExtractionError(
        f"{Path(filename).suffix or 'This file type'} is not supported."
    )
